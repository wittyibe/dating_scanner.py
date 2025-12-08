import requests
from bs4 import BeautifulSoup
import re
import time
import asyncio
import aiohttp
from typing import List, Dict
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import undetected_chromedriver as uc
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class ComprehensiveDatingScanner:
    def __init__(self):
        self.results = {}
        self.comprehensive_sites = {
            # Major Dating Sites (70+)
            'tinder.com': {'path': '/app/login', 'type': 'dating'},
            'bumble.com': {'path': '/login', 'type': 'dating'},
            'okcupid.com': {'path': '/login', 'type': 'dating'},
            'match.com': {'path': '/login', 'type': 'dating'},
            'plentyoffish.com': {'path': '/inbox', 'type': 'dating'},
            'eharmony.com': {'path': '/login', 'type': 'dating'},
            'eliteシングルズ.com': {'path': '/login', 'type': 'dating'},
            'zoosk.com': {'path': '/login', 'type': 'dating'},
            'chemistry.com': {'path': '/login', 'type': 'dating'},
            'silversingles.com': {'path': '/login', 'type': 'dating'},
            
            # Adult/Sex Sites
            'adultfriendfinder.com': {'path': '/login.php', 'type': 'adult'},
            'adultfriendfinder.net': {'path': '/login', 'type': 'adult'},
            'ashleymadison.com': {'path': '/login', 'type': 'adult'},
            'fetlife.com': {'path': '/login', 'type': 'fetish'},
            'alt.com': {'path': '/login', 'type': 'bdsm'},
            'bondage.com': {'path': '/login', 'type': 'bdsm'},
            'bdsm.com': {'path': '/login', 'type': 'bdsm'},
            ' Collarme.com': {'path': '/login', 'type': 'bdsm'},
            
            # Hookup Apps
            'pure.app': {'path': '/login', 'type': 'hookup'},
            'feeld.co': {'path': '/login', 'type': 'hookup'},
            'grindr.com': {'path': '/login', 'type': 'gay'},
            'scruff.com': {'path': '/login', 'type': 'gay'},
            'hornet.com': {'path': '/login', 'type': 'gay'},
            'romeo.com': {'path': '/login', 'type': 'gay'},
            'jackd.com': {'path': '/login', 'type': 'gay'},
            'growlr.com': {'path': '/login', 'type': 'gay'},
            
            # International
            'badoo.com': {'path': '/login', 'type': 'dating'},
            'happn.com': {'path': '/login', 'type': 'dating'},
            'hily.com': {'path': '/login', 'type': 'dating'},
            'innercircle.co': {'path': '/login', 'type': 'dating'},
            'theleague.com': {'path': '/login', 'type': 'dating'},
            'hinge.co': {'path': '/login', 'type': 'dating'},
            
            # Niche Adult
            'passion.com': {'path': '/login', 'type': 'adult'},
            'getiton.com': {'path': '/login', 'type': 'hookup'},
            'ihookup.com': {'path': '/login', 'type': 'hookup'},
            'xmatch.com': {'path': '/login', 'type': 'adult'},
            'together2night.com': {'path': '/login', 'type': 'hookup'},
            'snapmilfs.com': {'path': '/login', 'type': 'adult'},
            'snapfuck.com': {'path': '/login', 'type': 'adult'},
            'affairdating.com': {'path': '/login', 'type': 'adult'},
            
            # More Fetish/BDSM
            'kink.com': {'path': '/login', 'type': 'fetish'},
            'recon.com': {'path': '/login', 'type': 'fetish'},
            'fetster.com': {'path': '/login', 'type': 'fetish'},
            'thecage.co': {'path': '/login', 'type': 'fetish'},
            
            # Swingers
            'swinglifestyle.com': {'path': '/login', 'type': 'swinger'},
            'sdcfriends.com': {'path': '/login', 'type': 'swinger'},
            'kasidie.com': {'path': '/login', 'type': 'swinger'},
            
            # Additional 30+ sites
            'adultfriendfinder.org': {'path': '/login', 'type': 'adult'},
            'cams.com': {'path': '/login', 'type': 'cam'},
            'chaturbate.com': {'path': '/login', 'type': 'cam'},
            'livejasmin.com': {'path': '/login', 'type': 'cam'},
            'stripchat.com': {'path': '/login', 'type': 'cam'},
            'myfreecams.com': {'path': '/login', 'type': 'cam'},
            'bongacams.com': {'path': '/login', 'type': 'cam'},
            'flirt4free.com': {'path': '/login', 'type': 'cam'},
            'imlive.com': {'path': '/login', 'type': 'cam'},
            'streamate.com': {'path': '/login', 'type': 'cam'},
            'onlyfans.com': {'path': '/login', 'type': 'content'},
            'fansly.com': {'path': '/login', 'type': 'content'},
            'pocketstars.com': {'path': '/login', 'type': 'content'},
            'loyal.fans': {'path': '/login', 'type': 'content'},
            'justfor.fans': {'path': '/login', 'type': 'content'},
            'erothots.co': {'path': '/login', 'type': 'content'},
            'localmilfselfies.com': {'path': '/login', 'type': 'adult'},
            'fdating.com': {'path': '/login', 'type': 'dating'},
            'mingle2.com': {'path': '/login', 'type': 'dating'},
            'speeddate.com': {'path': '/login', 'type': 'dating'}
        }
        
    def stealth_browser(self):
        options = uc.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
        driver = uc.Chrome(options=options)
        driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        return driver
    
    async def scan_dating_sites(self, email: str, phone: str) -> List[Dict]:
        """Scan all 70+ dating/sex sites"""
        results = []
        driver = self.stealth_browser()
        
        for site, config in self.comprehensive_sites.items():
            try:
                print(f"Scanning {site}...")
                url = f"https://{site}{config['path']}"
                driver.get(url)
                time.sleep(2)
                
                # Try forgot password/reset flows
                forgot_selectors = [
                    "//*[contains(text(), 'forgot') or contains(text(), 'recover') or contains(text(), 'reset')]",
                    "//a[contains(@href, 'reset') or contains(@href, 'recover')]",
                    "//*[contains(@id, 'forgot') or contains(@id, 'reset')]"
                ]
                
                forgot_clicked = False
                for selector in forgot_selectors:
                    try:
                        forgot_btn = WebDriverWait(driver, 3).until(
                            EC.element_to_be_clickable((By.XPATH, selector))
                        )
                        forgot_btn.click()
                        forgot_clicked = True
                        time.sleep(1)
                        break
                    except:
                        continue
                
                # Input email/phone
                input_selectors = [
                    "//input[@type='email']",
                    "//input[contains(@name,'email') or contains(@name,'username')]",
                    "//input[@placeholder*='email']"
                ]
                
                for selector in input_selectors:
                    try:
                        email_field = driver.find_element(By.XPATH, selector)
                        email_field.clear()
                        email_field.send_keys(email)
                        break
                    except:
                        continue
                
                # Submit
                submit_selectors = [
                    "//button[@type='submit']",
                    "//input[@type='submit']",
                    "//button[contains(text(), 'send') or contains(text(), 'Send')]"
                ]
                
                for selector in submit_selectors:
                    try:
                        submit_btn = driver.find_element(By.XPATH, selector)
                        submit_btn.click()
                        time.sleep(3)
                        break
                    except:
                        pass
                
                page_source = driver.page_source.lower()
                
                # Account existence indicators
                exists_indicators = [
                    'sent to your email', 'check your inbox', 'reset link sent',
                    'account found', 'password reset', 'last login', 'check your email',
                    'we found your account', 'recovery email sent'
                ]
                
                no_account_indicators = [
                    'no account found', 'not registered', 'email not found',
                    'no user with that email', 'account does not exist'
                ]
                
                if any(indicator in page_source for indicator in exists_indicators):
                    results.append({
                        'site': site,
                        'type': config['type'],
                        'status': 'ACCOUNT_EXISTS',
                        'confidence': 'HIGH',
                        'scan_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    })
                elif any(indicator in page_source for indicator in no_account_indicators):
                    results.append({
                        'site': site,
                        'type': config['type'],
                        'status': 'NO_ACCOUNT',
                        'confidence': 'HIGH',
                        'scan_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    })
                else:
                    results.append({
                        'site': site,
                        'type': config['type'],
                        'status': 'UNKNOWN',
                        'confidence': 'LOW',
                        'scan_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    })
                    
            except Exception as e:
                results.append({
                    'site': site,
                    'type': config['type'],
                    'status': 'SCAN_FAILED',
                    'error': str(e)[:100],
                    'scan_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                })
            
            time.sleep(1)  # Rate limiting
        
        driver.quit()
        return results
    
    def generate_doc_report(self, email: str, phone: str, findings: List[Dict]) -> str:
        """Generate comprehensive DOCX report"""
        doc = Document()
        doc.add_heading(f'Dating/Sex Site Account Scan Report', 0)
        doc.add_heading(f'Target: {email} | {phone}', level=1)
        doc.add_paragraph(f'Scan completed: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total sites scanned: {len(self.comprehensive_sites)}')
        
        # Summary table
        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.style = 'Table Grid'
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Status'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Sites'
        hdr_cells[3].text = '%'
        
        exists_count = len([f for f in findings if f['status'] == 'ACCOUNT_EXISTS'])
        total_scanned = len([f for f in findings if f['status'] != 'SCAN_FAILED'])
        
        row_cells = summary_table.add_row().cells
        row_cells[0].text = 'ACCOUNTS FOUND'
        row_cells[1].text = str(exists_count)
        row_cells[2].text = ', '.join([f['site'] for f in findings if f['status'] == 'ACCOUNT_EXISTS'][:10])
        row_cells[3].text = f'{exists_count/total_scanned*100:.1f}%' if total_scanned else '0%'
        
        # Detailed findings by category
        categories = {}
        for finding in findings:
            cat = finding['type'].upper()
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(finding)
        
        for category, sites in categories.items():
            doc.add_heading(category, level=1)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Site'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Confidence'
            hdr_cells[3].text = 'Scan Time'
            hdr_cells[4].text = 'Notes'
            
            exists_sites = [s for s in sites if s['status'] == 'ACCOUNT_EXISTS']
            for site in exists_sites:
                row_cells = table.add_row().cells
                row_cells[0].text = site['site']
                row_cells[1].text = site['status']
                row_cells[2].text = site['confidence']
                row_cells[3].text = site['scan_time']
                row_cells[4].text = '⚠️ POTENTIAL ACCOUNT'
            
            if not exists_sites:
                p = table.add_row().cells[0].paragraph
                p.text = 'No accounts detected'
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        filename = f"dating_scan_{email.replace('@', '_at_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        return filename

async def main():
    email = input("Enter email to scan: ")
    phone = input("Enter phone (optional): ")
    
    scanner = ComprehensiveDatingScanner()
    print("🚀 Starting comprehensive scan of 70+ dating/sex sites...")
    
    findings = await scanner.scan_dating_sites(email, phone)
    
    # Generate DOCX report
    doc_file = scanner.generate_doc_report(email, phone, findings)
    print(f"\n✅ SCAN COMPLETE!")
    print(f"📄 Report saved: {doc_file}")
    print(f"🔍 ACCOUNTS FOUND: {len([f for f in findings if f['status']=='ACCOUNT_EXISTS'])}")
    print("\nDouble-click the DOCX file to view full results!")

if __name__ == "__main__":
    asyncio.run(main())

